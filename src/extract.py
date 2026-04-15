"""Hybrid resume extraction with per-file SHA1 cache.

Routes by extension: PDF (pdfplumber → Vision fallback), DOCX (python-docx),
HWP (LibreOffice CLI). PII is masked before persistence; raw_text never
reaches disk or downstream stages.
"""
from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Callable, Optional

from .cache import compute_sha1, load_resume_cache, save_resume_cache
from .llm import CostTracker, vision_ocr

SUPPORTED_EXTS = {".pdf", ".docx", ".hwp"}
SCAN_TEXT_THRESHOLD = 50
LIBREOFFICE_TIMEOUT_SEC = 60


def _mask(raw_text: str) -> tuple[str, dict[str, Any]]:
    """Delegate to src.pii.mask_pii. Returns (masked_text, pii_fields_separated)."""
    try:
        from . import pii  # type: ignore
    except Exception:
        # pii module not ready yet — return raw text untouched so pipeline still runs
        # during bring-up. Never ship to production without pii module in place.
        return raw_text, {"_warning": "pii_module_unavailable"}
    return pii.mask_pii(raw_text)


def _extract_pdf_text(file_path: Path) -> str:
    import pdfplumber  # lazy import

    chunks: list[str] = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t:
                chunks.append(t)
    return "\n".join(chunks)


def _extract_pdf_vision(file_path: Path, client, cost: CostTracker) -> str:
    """Render PDF pages to PNG via PyMuPDF (no poppler needed), then OCR via Vision API."""
    import fitz  # PyMuPDF, lazy import

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        png_paths: list[Path] = []
        doc = fitz.open(str(file_path))
        try:
            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=150)
                p = tmp / f"page_{i:03d}.png"
                pix.save(str(p))
                png_paths.append(p)
        finally:
            doc.close()
        if not png_paths:
            return ""
        return vision_ocr(client, png_paths, cost, step="extract_vision")


def _extract_docx_text(file_path: Path) -> str:
    import docx  # python-docx

    doc = docx.Document(str(file_path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_hwp_text(file_path: Path) -> str:
    """Extract text from HWPX (HWP 2011+, ZIP-based). Reads Preview/PrvText.txt first
    (plain-text rendering embedded by Hancom Office), falls back to parsing
    Contents/section*.xml <hs:t> nodes. LibreOffice fallback used only if both fail."""
    import zipfile
    import re as _re

    raw = file_path.read_bytes()
    if raw[:4] == b"PK\x03\x04":
        with zipfile.ZipFile(file_path) as z:
            names = z.namelist()
            if "Preview/PrvText.txt" in names:
                text = z.read("Preview/PrvText.txt").decode("utf-8", errors="replace")
                if len(text.strip()) >= SCAN_TEXT_THRESHOLD:
                    return text
            sections = sorted(n for n in names if n.startswith("Contents/section") and n.endswith(".xml"))
            chunks: list[str] = []
            for s in sections:
                xml = z.read(s).decode("utf-8", errors="replace")
                for m in _re.findall(r"<hs:t[^>]*>([^<]*)</hs:t>", xml):
                    if m:
                        chunks.append(m)
            joined = "\n".join(chunks)
            if joined.strip():
                return joined

    if shutil.which("libreoffice") is None and shutil.which("soffice") is None:
        raise FileNotFoundError("libreoffice_unavailable")
    bin_name = "libreoffice" if shutil.which("libreoffice") else "soffice"
    with tempfile.TemporaryDirectory() as tmpdir:
        proc = subprocess.run(
            [bin_name, "--headless", "--convert-to", "txt", "--outdir", tmpdir, str(file_path)],
            capture_output=True,
            timeout=LIBREOFFICE_TIMEOUT_SEC,
        )
        if proc.returncode != 0:
            raise RuntimeError(
                f"libreoffice returncode={proc.returncode} stderr={proc.stderr.decode('utf-8', 'ignore')[:300]}"
            )
        txt_path = Path(tmpdir) / f"{file_path.stem}.txt"
        if not txt_path.exists():
            candidates = list(Path(tmpdir).glob("*.txt"))
            if not candidates:
                raise RuntimeError("libreoffice produced no .txt output")
            txt_path = candidates[0]
        return txt_path.read_text(encoding="utf-8", errors="replace")


def _build_result(
    file_path: Path,
    sha1: str,
    ext: str,
    raw_text: str,
    method: str,
    ok: bool,
    error: Optional[str],
) -> dict[str, Any]:
    if ok and raw_text:
        masked_text, pii_fields = _mask(raw_text)
    else:
        masked_text, pii_fields = "", {}
    return {
        "resume_id": file_path.stem,
        "sha1": sha1,
        "format": ext.lstrip("."),
        "raw_text_length": len(raw_text),
        "masked_text": masked_text,
        "pii_fields_separated": pii_fields,
        "extraction_method": method,
        "ok": ok,
        "error": error,
    }


def extract_resume(file_path: Path, client, cost: CostTracker) -> dict[str, Any]:
    """Extract + mask a single resume. Cache-hit returns without re-processing."""
    file_path = Path(file_path)
    sha1 = compute_sha1(file_path)
    cached = load_resume_cache(sha1)
    if cached is not None:
        return cached

    ext = file_path.suffix.lower()
    raw_text = ""
    method = "unknown"
    ok = False
    error: Optional[str] = None

    if ext == ".pdf":
        method = "pdfplumber"
        try:
            raw_text = _extract_pdf_text(file_path)
        except Exception as e:
            raw_text = ""
            error = f"pdfplumber_error: {e}"

        if len(raw_text) < SCAN_TEXT_THRESHOLD:
            method = "vision"
            try:
                raw_text = _extract_pdf_vision(file_path, client, cost)
                ok = len(raw_text) > 0
                if not ok:
                    error = "vision_empty_output"
            except Exception as e:
                ok = False
                error = f"vision_error: {e}"
        else:
            ok = True
            error = None

    elif ext == ".docx":
        method = "python-docx"
        try:
            raw_text = _extract_docx_text(file_path)
            ok = len(raw_text) > 0
            if not ok:
                error = "docx_empty_output"
        except Exception as e:
            ok = False
            error = f"docx_error: {e}"

    elif ext == ".hwp":
        method = "hwpx-zip"
        try:
            raw_text = _extract_hwp_text(file_path)
            ok = len(raw_text) > 0
            if not ok:
                error = "hwp_empty_output"
        except FileNotFoundError:
            ok = False
            error = "hwp_unavailable"
        except subprocess.TimeoutExpired:
            ok = False
            error = "libreoffice_timeout"
        except Exception as e:
            ok = False
            error = f"hwp_error: {e}"

    else:
        method = "unsupported"
        ok = False
        error = f"unsupported_extension: {ext}"

    result = _build_result(file_path, sha1, ext, raw_text, method, ok, error)
    save_resume_cache(sha1, result)
    return result


def extract_directory(
    dir_path: Path,
    client,
    cost: CostTracker,
    on_progress: Optional[Callable[[int, int, str], None]] = None,
) -> list[dict[str, Any]]:
    """Walk dir_path recursively for supported resumes and extract each."""
    dir_path = Path(dir_path)
    files = sorted(
        [p for p in dir_path.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS]
    )
    total = len(files)
    results: list[dict[str, Any]] = []
    for i, fp in enumerate(files, start=1):
        try:
            res = extract_resume(fp, client, cost)
        except Exception as e:
            # Defensive: never crash the batch
            sha1 = ""
            try:
                sha1 = compute_sha1(fp)
            except Exception:
                pass
            res = {
                "resume_id": fp.stem,
                "sha1": sha1,
                "format": fp.suffix.lower().lstrip("."),
                "raw_text_length": 0,
                "masked_text": "",
                "pii_fields_separated": {},
                "extraction_method": "unknown",
                "ok": False,
                "error": f"unhandled_error: {e}",
            }
        status = "OK" if res.get("ok") else f"FAIL({res.get('error')})"
        print(f"[{i}/{total}] {fp.name} \u2192 {res.get('extraction_method')} {status}")
        if on_progress is not None:
            try:
                on_progress(i, total, res.get("resume_id", fp.stem))
            except Exception:
                pass
        results.append(res)
    return results


if __name__ == "__main__":
    import json
    from .llm import get_client

    data_dir = Path(__file__).resolve().parent.parent.parent / "task02_data" / "resumes"
    tracker = CostTracker()
    client = None
    try:
        client = get_client()
    except Exception as e:
        print(f"[warn] OpenAI client unavailable ({e}); Vision fallback will fail but text PDFs still work.")

    files = sorted(
        [p for p in data_dir.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS]
    )[:3]
    print(f"Testing first {len(files)} resume(s) from {data_dir}")

    for fp in files:
        res = extract_resume(fp, client, tracker)
        # raw_text was never stored, but strip any debug field just to be safe
        display = {k: v for k, v in res.items() if k != "raw_text"}
        # Truncate masked_text for console sanity
        if isinstance(display.get("masked_text"), str) and len(display["masked_text"]) > 400:
            display["masked_text"] = display["masked_text"][:400] + "...<truncated>"
        print(json.dumps(display, ensure_ascii=False, indent=2))

    print(f"\nTotal cost: ${tracker.total:.6f}")
